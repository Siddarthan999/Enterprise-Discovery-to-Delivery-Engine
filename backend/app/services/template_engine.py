import re
import subprocess
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from io import BytesIO


# ----------------------------
# FONT AVAILABILITY (for PDF path only)
# ----------------------------
_INSTALLED_FONTS_CACHE = None

# Metric-compatible substitutes for common MS-style fonts, chosen because
# they're widely available on Linux (via fonts-liberation / fonts-crosextra
# packages) and match glyph widths closely enough that line breaks/layout
# don't shift. Extend this table as you discover more gaps in your
# specific server environment.
_FONT_SUBSTITUTES = {
    "calibri": "Carlito",
    "cambria": "Caladea",
    "arial": "Liberation Sans",
    "times new roman": "Liberation Serif",
    "courier new": "Liberation Mono",
    "georgia": "Liberation Serif",
    "verdana": "DejaVu Sans",
    "helvetica": "Liberation Sans",
}

_GENERIC_FALLBACK_SERIF = "Liberation Serif"
_GENERIC_FALLBACK_SANS = "Liberation Sans"


def _get_installed_font_names() -> set:
    global _INSTALLED_FONTS_CACHE
    if _INSTALLED_FONTS_CACHE is not None:
        return _INSTALLED_FONTS_CACHE

    try:
        result = subprocess.run(
            ["fc-list", ":", "family"], capture_output=True, text=True, timeout=5
        )
        names = set()
        for line in result.stdout.splitlines():
            for name in line.split(","):
                names.add(name.strip().lower())
        _INSTALLED_FONTS_CACHE = names
    except Exception:
        # fontconfig unavailable — don't crash export, just skip substitution.
        _INSTALLED_FONTS_CACHE = set()

    return _INSTALLED_FONTS_CACHE


def _resolve_font_for_pdf(font_name: str) -> str:
    """Only used when building the DOCX that will be converted to PDF via
    LibreOffice. If the requested font isn't actually installed on this
    server, LibreOffice would otherwise silently substitute something
    unpredictable — instead substitute deliberately to a known-good,
    metric-compatible font."""
    if not font_name:
        return _GENERIC_FALLBACK_SANS

    installed = _get_installed_font_names()
    if font_name.lower() in installed:
        return font_name

    substitute = _FONT_SUBSTITUTES.get(font_name.lower())
    if substitute and substitute.lower() in installed:
        return substitute

    # Last resort: guess serif vs sans from the name and use whichever
    # generic fallback is actually installed.
    serif_hint = any(k in font_name.lower() for k in ["times", "georgia", "cambria", "serif", "garamond"])
    fallback = _GENERIC_FALLBACK_SERIF if serif_hint else _GENERIC_FALLBACK_SANS
    return fallback if fallback.lower() in installed else font_name


# ----------------------------
# FONT DETECTION
# ----------------------------
def _detect_brand_profile(doc: Document) -> dict:
    largest_run = None
    body_font_name = None
    body_color = None

    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text or not run.text.strip():
                continue

            size = run.font.size.pt if run.font.size else None
            if size and (largest_run is None or size > largest_run[0]):
                largest_run = (size, run)

            if run.font.name and (size is None or size <= 14) and not body_font_name:
                body_font_name = run.font.name
                if run.font.color and run.font.color.rgb:
                    body_color = run.font.color.rgb

    normal_style_font = doc.styles["Normal"].font
    if not body_font_name:
        body_font_name = normal_style_font.name or "Calibri"
    if not body_color and normal_style_font.color and normal_style_font.color.rgb:
        body_color = normal_style_font.color.rgb

    if largest_run:
        title_size, title_run = largest_run
        title_font_name = title_run.font.name or body_font_name
        title_bold = bool(title_run.font.bold)
        title_color = None
        if title_run.font.color and title_run.font.color.rgb:
            title_color = title_run.font.color.rgb
    else:
        title_size = 28
        title_font_name = body_font_name
        title_bold = True
        title_color = None

    return {
        "title_font": title_font_name,
        "title_size": title_size,
        "title_bold": title_bold,
        "title_color": title_color,
        "body_font": body_font_name,
        "body_color": body_color,
    }


def _get_or_create_style(doc: Document, name: str, font_name: str, size_pt: float,
                          bold: bool = False, color: RGBColor | None = None,
                          base_style: str | None = None):
    if name in [s.name for s in doc.styles]:
        return doc.styles[name]

    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    if base_style and base_style in [s.name for s in doc.styles]:
        style.base_style = doc.styles[base_style]

    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color:
        style.font.color.rgb = color

    return style


def _build_sow_styles(doc: Document, for_pdf: bool = False) -> dict:
    profile = _detect_brand_profile(doc)

    title_font = profile["title_font"]
    body_font = profile["body_font"]

    if for_pdf:
        title_font = _resolve_font_for_pdf(title_font)
        body_font = _resolve_font_for_pdf(body_font)

    title_size = profile["title_size"]
    h1_size = max(18, round(title_size * 0.55))
    h2_size = max(14, round(title_size * 0.4))
    body_size = 11

    styles = {}

    styles["title"] = _get_or_create_style(
        doc, "SOW Title", title_font, title_size,
        bold=True, color=profile["title_color"]
    )
    styles["h1"] = _get_or_create_style(
        doc, "SOW Heading 1", title_font, h1_size,
        bold=True, color=profile["title_color"]
    )
    styles["h2"] = _get_or_create_style(
        doc, "SOW Heading 2", title_font, h2_size,
        bold=profile["title_bold"], color=profile["body_color"]
    )
    styles["body"] = _get_or_create_style(
        doc, "SOW Body", body_font, body_size,
        bold=False, color=profile["body_color"]
    )

    has_list_style = "List Bullet" in [s.name for s in doc.styles]
    styles["bullet"] = _get_or_create_style(
        doc, "SOW Bullet", body_font, body_size,
        bold=False, color=profile["body_color"],
        base_style="List Bullet" if has_list_style else None,
    )
    styles["_has_native_bullet"] = has_list_style

    return styles


# ----------------------------
# COVER PAGE PLACEHOLDER REPLACEMENT
# ----------------------------

# Exact (case-insensitive, whitespace-trimmed) paragraph/content-control text
# this engine recognizes as a placeholder label, mapped to the cover_fields
# key that should replace it. Extend this as you see more template
# conventions. Includes both the "clean label" form (e.g. "Company Name")
# and the bracketed "instructional placeholder" form some cover-page
# templates use (e.g. "[Company Name]" or Word's own
# "[Type the company name]" ghost text) — add more bracketed variants here
# if your diagnostic print turns up different wording.
_LITERAL_PLACEHOLDER_LABELS = {
    "company name": "company_name",
    "sub-headline": "sub_headline",
    "subheadline": "sub_headline",
    "[company name]": "company_name",
    "[sub-headline]": "sub_headline",
    "[type the company name]": "company_name",
    "[type the document subtitle]": "sub_headline",
}

# Content controls are sometimes tagged (w:tag) or aliased (w:alias) by
# whoever built the template. Matching on tag/alias is more robust than
# matching on displayed text, since it survives edits to the visible
# placeholder wording. Add mappings here if you discover tags in your
# template (check via the diagnostic snippet mentioned below).
_SDT_TAG_FIELD_MAP = {
    "companyname": "company_name",
    "company_name": "company_name",
    "subheadline": "sub_headline",
    "sub_headline": "sub_headline",
    "sub-headline": "sub_headline",
}

_COMPLETED_ON_PREFIX = "this statement of work completed on"


def _set_paragraph_text_preserve_format(paragraph, new_text: str):
    """Replaces a paragraph's visible text while keeping the first run's
    font formatting (name/size/bold/color) — critical so cover-page
    styling survives the substitution instead of reverting to default."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _sdt_placeholder_text(sdt_element) -> str:
    """Concatenate all w:t text inside one content control (SDT)."""
    return "".join(t.text or "" for t in sdt_element.iter(qn("w:t")))


def _sdt_tag_or_alias(sdt_element) -> str | None:
    """Return the w:tag value, falling back to w:alias, lowercased."""
    sdt_pr = sdt_element.find(qn("w:sdtPr"))
    if sdt_pr is None:
        return None

    tag_el = sdt_pr.find(qn("w:tag"))
    if tag_el is not None:
        val = tag_el.get(qn("w:val"))
        if val:
            return val.strip().lower()

    alias_el = sdt_pr.find(qn("w:alias"))
    if alias_el is not None:
        val = alias_el.get(qn("w:val"))
        if val:
            return val.strip().lower()

    return None


def _set_sdt_text(sdt_element, new_text: str):
    """Content controls can have multiple w:t runs. Put the new text in
    the first w:t and blank the rest, same pattern as
    _set_paragraph_text_preserve_format so the first run's formatting
    (and thus the control's font/color) is preserved."""
    t_elements = list(sdt_element.iter(qn("w:t")))
    if not t_elements:
        return
    t_elements[0].text = new_text
    for t in t_elements[1:]:
        t.text = ""


def _resolve_field_value(field_key: str, cover_fields: dict, completion_date: str) -> str | None:
    if field_key == "completion_date":
        return completion_date
    return cover_fields.get(field_key)


def _apply_token_syntax(raw_text: str, cover_fields: dict, completion_date: str) -> str | None:
    """Returns the substituted text if {{TOKEN}} syntax was found and
    changed something, else None."""
    if "{{" not in raw_text or "}}" not in raw_text:
        return None

    def _replace(match):
        key = match.group(1).strip().lower()
        value = _resolve_field_value(key, cover_fields, completion_date)
        return str(value) if value is not None else ""

    new_text = re.sub(r"\{\{\s*([\w_]+)\s*\}\}", _replace, raw_text)
    return new_text if new_text != raw_text else None


def apply_cover_page_fields(doc: Document, cover_fields: dict | None):
    """Three mechanisms, all scoped to the template's ORIGINAL content
    only (called before any generated content is added, so there's no
    risk of matching against the SOW body text):

    1. {{TOKEN}} syntax — explicit, safe to match anywhere. Recommended
       for new templates going forward. Checked in both plain paragraphs
       and content controls.
    2. Content-control (w:sdt) tag/alias matching — the most robust way
       to target cover-page fields when the template was built with
       Word's "Insert Cover Page" gallery or custom content controls,
       since python-docx's doc.paragraphs does NOT walk into w:sdt
       blocks at all.
    3. Known literal placeholder labels (e.g. text that says exactly
       "Company Name" or "[Type the company name]") — supports both
       plain-paragraph templates and content controls that have no
       tag/alias set, by matching on displayed text as a fallback.
    """
    if not cover_fields:
        cover_fields = {}

    completion_date = cover_fields.get(
        "completion_date", datetime.utcnow().strftime("%B %d, %Y")
    )

    # --- Pass 1: plain paragraphs (doc.paragraphs) ---
    for para in doc.paragraphs:
        raw_text = para.text
        stripped = raw_text.strip()
        lowered = stripped.lower()

        new_text = _apply_token_syntax(raw_text, cover_fields, completion_date)
        if new_text is not None:
            _set_paragraph_text_preserve_format(para, new_text)
            continue

        if lowered in _LITERAL_PLACEHOLDER_LABELS:
            field_key = _LITERAL_PLACEHOLDER_LABELS[lowered]
            value = cover_fields.get(field_key, "")
            if value:
                _set_paragraph_text_preserve_format(para, value)
            continue

        if lowered.startswith(_COMPLETED_ON_PREFIX):
            _set_paragraph_text_preserve_format(
                para, f"This statement of work completed on {completion_date}"
            )

    # --- Pass 2: content controls (w:sdt) ---
    # These are invisible to doc.paragraphs, so they need direct XML
    # traversal. Common in Word's built-in cover-page templates (e.g.
    # "Facet", "Ion", "Motion") where "Company Name" / subtitle fields
    # are implemented as structured document tags, not plain text.
    for sdt in doc.element.body.iter(qn("w:sdt")):
        raw_text = _sdt_placeholder_text(sdt)
        if not raw_text:
            continue

        stripped = raw_text.strip()
        lowered = stripped.lower()

        # 2a. Token syntax inside a content control
        new_text = _apply_token_syntax(raw_text, cover_fields, completion_date)
        if new_text is not None:
            _set_sdt_text(sdt, new_text)
            continue

        # 2b. Tag/alias-based matching (most robust)
        tag = _sdt_tag_or_alias(sdt)
        if tag and tag in _SDT_TAG_FIELD_MAP:
            field_key = _SDT_TAG_FIELD_MAP[tag]
            value = cover_fields.get(field_key, "")
            if value:
                _set_sdt_text(sdt, value)
            continue

        # 2c. Literal displayed-text fallback
        if lowered in _LITERAL_PLACEHOLDER_LABELS:
            field_key = _LITERAL_PLACEHOLDER_LABELS[lowered]
            value = cover_fields.get(field_key, "")
            if value:
                _set_sdt_text(sdt, value)
            continue

        if lowered.startswith(_COMPLETED_ON_PREFIX):
            _set_sdt_text(sdt, f"This statement of work completed on {completion_date}")


# ----------------------------
# TEMPLATE LOADER
# ----------------------------
def load_template(template_path: str) -> Document:
    return Document(template_path)


# ----------------------------
# ADD CONTENT TO DOCX
# ----------------------------
def add_sow_content(doc: Document, markdown: str, for_pdf: bool = False):

    styles = _build_sow_styles(doc, for_pdf=for_pdf)

    doc.add_page_break()

    lines = markdown.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            doc.add_paragraph(line[2:], style=styles["title"])

        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style=styles["h1"])

        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style=styles["h2"])

        elif line.startswith("- "):
            text = line[2:]
            if not styles["_has_native_bullet"]:
                text = f"• {text}"
            doc.add_paragraph(text, style=styles["bullet"])

        else:
            doc.add_paragraph(line, style=styles["body"])

    return doc


# ----------------------------
# MAIN EXPORT FUNCTION
# ----------------------------
def export_docx_from_template(template_path: str, sow_markdown: str,
                               cover_fields: dict | None = None,
                               for_pdf: bool = False) -> bytes:

    doc = load_template(template_path)

    # Must happen BEFORE add_sow_content — it only touches the template's
    # original paragraphs/content controls, and add_sow_content appends
    # new paragraphs after a page break, so ordering keeps this safely
    # scoped to the cover page.
    apply_cover_page_fields(doc, cover_fields)

    doc = add_sow_content(doc, sow_markdown, for_pdf=for_pdf)

    buffer = BytesIO()
    doc.save(buffer)

    return buffer.getvalue()