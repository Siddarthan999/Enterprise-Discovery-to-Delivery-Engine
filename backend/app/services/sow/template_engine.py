import re
import subprocess
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from io import BytesIO

# ----------------------------
# FONT AVAILABILITY (for PDF path only)
# ----------------------------
_INSTALLED_FONTS_CACHE = None

_FONT_SUBSTITUTES = {
    "calibri": "Carlito",
    "cambria": "Caladea",
    "arial": "Liberation Sans",
    "times new roman": "Liberation Serif",
    "courier new": "Liberation Mono",
    "georgia": "Liberation Serif",
    "verdana": "DejaVu Sans",
    "helvetica": "Liberation Sans",
}

_GENERIC_FALLBACK_SERIF = "Liberation Serif"
_GENERIC_FALLBACK_SANS = "Liberation Sans"
_SOW_BODY_PLACEHOLDER = "{{SOW_CONTENT}}"


def _get_installed_font_names() -> set:
    global _INSTALLED_FONTS_CACHE
    if _INSTALLED_FONTS_CACHE is not None:
        return _INSTALLED_FONTS_CACHE

    try:
        result = subprocess.run(
            ["fc-list", ":", "family"], capture_output=True, text=True, timeout=5
        )
        names = set()
        for line in result.stdout.splitlines():
            for name in line.split(","):
                names.add(name.strip().lower())
        _INSTALLED_FONTS_CACHE = names
    except Exception:
        _INSTALLED_FONTS_CACHE = set()

    return _INSTALLED_FONTS_CACHE


def _resolve_font_for_pdf(font_name: str) -> str:
    if not font_name:
        return _GENERIC_FALLBACK_SANS

    installed = _get_installed_font_names()
    if font_name.lower() in installed:
        return font_name

    substitute = _FONT_SUBSTITUTES.get(font_name.lower())
    if substitute and substitute.lower() in installed:
        return substitute

    serif_hint = any(k in font_name.lower() for k in ["times", "georgia", "cambria", "serif", "garamond"])
    fallback = _GENERIC_FALLBACK_SERIF if serif_hint else _GENERIC_FALLBACK_SANS
    return fallback if fallback.lower() in installed else font_name


# ----------------------------
# FONT DETECTION
# ----------------------------
def _detect_brand_profile(doc: Document) -> dict:
    largest_run = None
    body_font_name = None
    body_color = None

    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text or not run.text.strip():
                continue

            size = run.font.size.pt if run.font.size else None
            if size and (largest_run is None or size > largest_run[0]):
                largest_run = (size, run)

            if run.font.name and (size is None or size <= 14) and not body_font_name:
                body_font_name = run.font.name
                if run.font.color and run.font.color.rgb:
                    body_color = run.font.color.rgb

    normal_style_font = doc.styles["Normal"].font
    if not body_font_name:
        body_font_name = normal_style_font.name or "Calibri"
    if not body_color and normal_style_font.color and normal_style_font.color.rgb:
        body_color = normal_style_font.color.rgb

    if largest_run:
        title_size, title_run = largest_run
        title_font_name = title_run.font.name or body_font_name
        title_bold = bool(title_run.font.bold)
        title_color = title_run.font.color.rgb if title_run.font.color and title_run.font.color.rgb else None
    else:
        title_size = 28
        title_font_name = body_font_name
        title_bold = True
        title_color = None

    return {
        "title_font": title_font_name,
        "title_size": title_size,
        "title_bold": title_bold,
        "title_color": title_color,
        "body_font": body_font_name,
        "body_color": body_color,
    }


def _get_or_create_style(doc: Document, name: str, font_name: str, size_pt: float,
                          bold: bool = False, color: RGBColor | None = None,
                          base_style: str | None = None,
                          space_before_pt: float = 0, space_after_pt: float = 0):
    if name in [s.name for s in doc.styles]:
        return doc.styles[name]

    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    if base_style and base_style in [s.name for s in doc.styles]:
        style.base_style = doc.styles[base_style]

    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color:
        style.font.color.rgb = color

    # Fix 2: give heading-level styles real spacing above/below so numbered
    # section headings don't render flush against surrounding body text.
    style.paragraph_format.space_before = Pt(space_before_pt)
    style.paragraph_format.space_after = Pt(space_after_pt)

    return style


def _build_sow_styles(doc: Document, for_pdf: bool = False) -> dict:
    profile = _detect_brand_profile(doc)

    title_font = profile["title_font"]
    body_font = profile["body_font"]

    if for_pdf:
        title_font = _resolve_font_for_pdf(title_font)
        body_font = _resolve_font_for_pdf(body_font)

    title_size = profile["title_size"]
    h1_size = max(18, round(title_size * 0.55))
    h2_size = max(14, round(title_size * 0.4))
    body_size = 11

    styles = {}

    styles["title"] = _get_or_create_style(
        doc, "SOW Title", title_font, title_size,
        bold=True, color=profile["title_color"],
        space_before_pt=0, space_after_pt=18,
    )
    styles["h1"] = _get_or_create_style(
        doc, "SOW Heading 1", title_font, h1_size,
        bold=True, color=profile["title_color"],
        space_before_pt=18, space_after_pt=8,
    )
    styles["h2"] = _get_or_create_style(
        doc, "SOW Heading 2", title_font, h2_size,
        bold=profile["title_bold"], color=profile["body_color"],
        space_before_pt=12, space_after_pt=6,
    )
    styles["body"] = _get_or_create_style(
        doc, "SOW Body", body_font, body_size,
        bold=False, color=profile["body_color"],
        space_before_pt=0, space_after_pt=0,
    )

    has_list_style = "List Bullet" in [s.name for s in doc.styles]
    styles["bullet"] = _get_or_create_style(
        doc, "SOW Bullet", body_font, body_size,
        bold=False, color=profile["body_color"],
        base_style="List Bullet" if has_list_style else None,
        space_before_pt=0, space_after_pt=0,
    )

    has_number_style = "List Number" in [s.name for s in doc.styles]
    styles["number"] = _get_or_create_style(
        doc, "SOW Number", body_font, body_size,
        bold=False, color=profile["body_color"],
        base_style="List Number" if has_number_style else None,
        space_before_pt=0, space_after_pt=0,
    )

    styles["_has_native_bullet"] = has_list_style
    styles["_has_native_number"] = has_number_style
    return styles


# ----------------------------
# COVER PAGE PLACEHOLDER REPLACEMENT
# ----------------------------
_LITERAL_PLACEHOLDER_LABELS = {
    "company name": "company_name",
    "sub-headline": "sub_headline",
    "subheadline": "sub_headline",
    "[company name]": "company_name",
    "[sub-headline]": "sub_headline",
    "[type the company name]": "company_name",
    "[type the document subtitle]": "sub_headline",
}

_SDT_TAG_FIELD_MAP = {
    "companyname": "company_name",
    "company_name": "company_name",
    "subheadline": "sub_headline",
    "sub_headline": "sub_headline",
    "sub-headline": "sub_headline",
}

_COMPLETED_ON_PREFIX = "this statement of work completed on"


def _set_paragraph_text_preserve_format(paragraph, new_text: str):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _remove_paragraph(paragraph):
    """Fix 1: fully remove a paragraph from the document body, rather than
    leaving a raw '[Company Name]' / 'Sub-Headline' placeholder visible when
    no value was supplied for that field."""
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _remove_sdt(sdt_element):
    """Fix 1 (content-control variant): remove a w:sdt node entirely when its
    mapped field has no value, instead of leaving the placeholder text in."""
    parent = sdt_element.getparent()
    if parent is not None:
        parent.remove(sdt_element)


def _sdt_placeholder_text(sdt_element) -> str:
    return "".join(t.text or "" for t in sdt_element.iter(qn("w:t")))


def _sdt_tag_or_alias(sdt_element) -> str | None:
    sdt_pr = sdt_element.find(qn("w:sdtPr"))
    if sdt_pr is None:
        return None

    tag_el = sdt_pr.find(qn("w:tag"))
    if tag_el is not None:
        val = tag_el.get(qn("w:val"))
        if val:
            return val.strip().lower()

    alias_el = sdt_pr.find(qn("w:alias"))
    if alias_el is not None:
        val = alias_el.get(qn("w:val"))
        if val:
            return val.strip().lower()

    return None


def _set_sdt_text(sdt_element, new_text: str):
    t_elements = list(sdt_element.iter(qn("w:t")))
    if not t_elements:
        return
    t_elements[0].text = new_text
    for t in t_elements[1:]:
        t.text = ""


def _resolve_field_value(field_key: str, cover_fields: dict, completion_date: str) -> str | None:
    if field_key == "completion_date":
        return completion_date
    return cover_fields.get(field_key)


def _apply_token_syntax(raw_text: str, cover_fields: dict, completion_date: str) -> str | None:
    if "{{" not in raw_text or "}}" not in raw_text:
        return None

    def _replace(match):
        key = match.group(1).strip().lower()
        value = _resolve_field_value(key, cover_fields, completion_date)
        return str(value) if value is not None else ""

    new_text = re.sub(r"\{\{\s*([\w_]+)\s*\}\}", _replace, raw_text)
    return new_text if new_text != raw_text else None


def apply_cover_page_fields(doc: Document, cover_fields: dict | None):
    if not cover_fields:
        cover_fields = {}

    completion_date = cover_fields.get(
        "completion_date", datetime.utcnow().strftime("%B %d, %Y")
    )

    # doc.paragraphs is already a materialized list, so it's safe to mutate
    # (remove paragraphs) while iterating it below.
    for para in doc.paragraphs:
        raw_text = para.text
        stripped = raw_text.strip()
        lowered = stripped.lower()

        new_text = _apply_token_syntax(raw_text, cover_fields, completion_date)
        if new_text is not None:
            _set_paragraph_text_preserve_format(para, new_text)
            continue

        if lowered in _LITERAL_PLACEHOLDER_LABELS:
            field_key = _LITERAL_PLACEHOLDER_LABELS[lowered]
            value = cover_fields.get(field_key, "")
            if value:
                _set_paragraph_text_preserve_format(para, value)
            else:
                # Fix 1: no value supplied -> remove the placeholder paragraph
                # entirely instead of leaving "Sub-Headline" / "Company Name"
                # sitting in the rendered document.
                _remove_paragraph(para)
            continue

        if lowered.startswith(_COMPLETED_ON_PREFIX):
            _set_paragraph_text_preserve_format(
                para, f"This statement of work completed on {completion_date}"
            )

    # doc.element.body.iter(qn("w:sdt")) is a *live* tree traversal, so it
    # must be materialized into a list before we start removing nodes from
    # the tree mid-iteration.
    for sdt in list(doc.element.body.iter(qn("w:sdt"))):
        raw_text = _sdt_placeholder_text(sdt)
        if not raw_text:
            continue

        stripped = raw_text.strip()
        lowered = stripped.lower()

        new_text = _apply_token_syntax(raw_text, cover_fields, completion_date)
        if new_text is not None:
            _set_sdt_text(sdt, new_text)
            continue

        tag = _sdt_tag_or_alias(sdt)
        if tag and tag in _SDT_TAG_FIELD_MAP:
            field_key = _SDT_TAG_FIELD_MAP[tag]
            value = cover_fields.get(field_key, "")
            if value:
                _set_sdt_text(sdt, value)
            else:
                _remove_sdt(sdt)
            continue

        if lowered in _LITERAL_PLACEHOLDER_LABELS:
            field_key = _LITERAL_PLACEHOLDER_LABELS[lowered]
            value = cover_fields.get(field_key, "")
            if value:
                _set_sdt_text(sdt, value)
            else:
                _remove_sdt(sdt)
            continue

        if lowered.startswith(_COMPLETED_ON_PREFIX):
            _set_sdt_text(sdt, f"This statement of work completed on {completion_date}")


# ----------------------------
# TEMPLATE LOADER
# ----------------------------
def load_template(template_path: str) -> Document:
    return Document(template_path)


# ----------------------------
# CONTENT INSERTION HELPERS
# ----------------------------
def _find_body_placeholder(doc: Document, placeholder: str = _SOW_BODY_PLACEHOLDER):
    for para in doc.paragraphs:
        if placeholder in (para.text or ""):
            return para
    return None


def _clear_placeholder_paragraph(paragraph):
    if paragraph is None:
        return
    if paragraph.runs:
        paragraph.runs[0].text = ""
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = ""


def _insert_paragraph_after(anchor: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = anchor.insert_paragraph_before(text=text, style=style)
    anchor._p.addnext(new_p._p)
    return new_p


def _add_inline_runs(paragraph, text: str):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _is_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _is_table_separator(line: str) -> bool:
    if not _is_table_line(line):
        return False
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def _split_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _insert_table_after(doc: Document, anchor: Paragraph, table_lines: list[str]) -> Paragraph:
    rows = [_split_table_row(line) for line in table_lines if _is_table_line(line)]
    if not rows:
        return anchor

    header = rows[0]
    body = rows[2:] if len(rows) > 1 and _is_table_separator(table_lines[1]) else rows[1:]

    table = doc.add_table(rows=1, cols=len(header))
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for idx, value in enumerate(header):
        cell_p = table.rows[0].cells[idx].paragraphs[0]
        _add_inline_runs(cell_p, value)
        for run in cell_p.runs:
            run.bold = True

    for row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row[:len(cells)]):
            cell_p = cells[idx].paragraphs[0]
            _add_inline_runs(cell_p, value)

    anchor._p.addnext(table._tbl)
    spacer = _insert_paragraph_after(anchor, "", style=None)
    table._tbl.addnext(spacer._p)
    return spacer


# ----------------------------
# ADD CONTENT TO DOCX
# ----------------------------
def add_sow_content(doc: Document, markdown: str, for_pdf: bool = False):
    styles = _build_sow_styles(doc, for_pdf=for_pdf)

    anchor = _find_body_placeholder(doc)
    use_placeholder = anchor is not None

    if use_placeholder:
        _clear_placeholder_paragraph(anchor)
        current = anchor
    else:
        doc.add_page_break()
        current = doc.paragraphs[-1]

    lines = markdown.splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i].rstrip("\n")
        line = raw.strip()

        if not line:
            i += 1
            continue

        if _is_table_line(line):
            table_lines = [line]
            i += 1
            while i < len(lines) and _is_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            current = _insert_table_after(doc, current, table_lines)
            continue

        if line.startswith("# "):
            p = _insert_paragraph_after(current, "", style=styles["title"])
            _add_inline_runs(p, line[2:].strip())
            current = p
            i += 1
            continue

        if line.startswith("## "):
            p = _insert_paragraph_after(current, "", style=styles["h1"])
            _add_inline_runs(p, line[3:].strip())
            current = p
            i += 1
            continue

        if line.startswith("### "):
            p = _insert_paragraph_after(current, "", style=styles["h2"])
            _add_inline_runs(p, line[4:].strip())
            current = p
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            style = styles["number"] if styles["_has_native_number"] else styles["body"]
            p = _insert_paragraph_after(current, "", style=style)
            if not styles["_has_native_number"]:
                text = f"{line}"
            _add_inline_runs(p, text)
            current = p
            i += 1
            continue

        if line.startswith("- "):
            text = line[2:].strip()
            p = _insert_paragraph_after(current, "", style=styles["bullet"])
            if not styles["_has_native_bullet"]:
                text = f"• {text}"
            _add_inline_runs(p, text)
            current = p
            i += 1
            continue

        p = _insert_paragraph_after(current, "", style=styles["body"])
        _add_inline_runs(p, line)
        current = p
        i += 1

    return doc


# ----------------------------
# MAIN EXPORT FUNCTION
# ----------------------------
def export_docx_from_template(template_path: str, sow_markdown: str,
                               cover_fields: dict | None = None,
                               for_pdf: bool = False) -> bytes:

    doc = load_template(template_path)
    apply_cover_page_fields(doc, cover_fields)
    doc = add_sow_content(doc, sow_markdown, for_pdf=for_pdf)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()