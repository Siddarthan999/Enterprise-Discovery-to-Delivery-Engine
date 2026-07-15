from docx import Document
from docx.oxml.ns import qn
import sys


def diagnose_template(template_path):
    doc = Document(template_path)

    print("=== 1. doc.paragraphs (top-level body paragraphs) ===")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"[{i}] {p.text!r}")

    print("\n=== 2. w:sdt content controls (anywhere in body) ===")
    for i, sdt in enumerate(doc.element.body.iter(qn('w:sdt'))):
        text = "".join(t.text or "" for t in sdt.iter(qn('w:t')))
        sdt_pr = sdt.find(qn('w:sdtPr'))
        tag = alias = None
        if sdt_pr is not None:
            tag_el = sdt_pr.find(qn('w:tag'))
            alias_el = sdt_pr.find(qn('w:alias'))
            tag = tag_el.get(qn('w:val')) if tag_el is not None else None
            alias = alias_el.get(qn('w:val')) if alias_el is not None else None
        if text.strip():
            print(f"[{i}] text={text!r} tag={tag!r} alias={alias!r}")

    print("\n=== 3. Table cell paragraphs ===")
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    if p.text.strip():
                        print(f"table[{ti}] row[{ri}] cell[{ci}]: {p.text!r}")

    print("\n=== 4. Text boxes (w:txbxContent) ===")
    for i, txbx in enumerate(doc.element.body.iter(qn('w:txbxContent'))):
        for p in txbx.iter(qn('w:p')):
            text = "".join(t.text or "" for t in p.iter(qn('w:t')))
            if text.strip():
                print(f"txbx[{i}]: {text!r}")

    print("\n=== 5. Raw document.xml dump (first 6000 chars) ===")
    print(doc.element.xml[:6000])


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python diagnose_template.py <path_to_template.docx>")
        sys.exit(1)

    diagnose_template(sys.argv[1])